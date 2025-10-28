import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import LabelEncoder, StandardScaler
import warnings
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

warnings.filterwarnings('ignore')

class BankMarketingAnalysis:
    def __init__(self, filepath):
        """Initialize the analysis with data from CSV file"""
        self.df = pd.read_csv(filepath)
        self.report = []
        self.analysis_history = []
        self.setup_data()
        
    def setup_data(self):
        """Prepare data for analysis"""
        self.le_dict = {}
        self.df_encoded = self.df.copy()
        
        categorical_cols = self.df_encoded.select_dtypes(include=['object']).columns
        for col in categorical_cols:
            le = LabelEncoder()
            self.df_encoded[col] = le.fit_transform(self.df_encoded[col])
            self.le_dict[col] = le
    
    def normalize_data(self):
        """
        TÍNH NĂNG / FEATURE: Chuẩn Hóa Dữ Liệu / Data Normalization
        MÔ TẢ / DESCRIPTION: Chuẩn hóa các biến định lượng bằng StandardScaler
        - Transform tất cả biến số về scale [mean=0, std=1]
        - Giúp các thuật toán ML hoạt động tốt hơn
        - Lưu scaler để transform dữ liệu mới
        """
        print("\n" + "="*80)
        print("CHUẨN HÓA DỮ LIỆU / DATA NORMALIZATION")
        print("="*80)
        
        # 1. Lấy các biến số
        numerical_cols = list(self.df.select_dtypes(include=[np.number]).columns)
        
        print(f"\n✓ Phát hiện {len(numerical_cols)} biến số cần chuẩn hóa:")
        print(f"  {numerical_cols}")
        
        # 2. Thống kê TRƯỚC chuẩn hóa
        print("\n" + "="*80)
        print("TRƯỚC CHUẨN HÓA / BEFORE NORMALIZATION")
        print("="*80)
        stats_before = self.df[numerical_cols].describe()
        print("\nThống kê mô tả:")
        print(stats_before)
        
        self.report.append("\n" + "="*80)
        self.report.append("CHUẨN HÓA DỮ LIỆU / DATA NORMALIZATION")
        self.report.append("="*80)
        self.report.append("\n" + "="*80)
        self.report.append("TRƯỚC CHUẨN HÓA / BEFORE NORMALIZATION")
        self.report.append("="*80)
        self.report.append("\nThống kê mô tả:")
        self.report.append(str(stats_before))
        
        # 3. Áp dụng StandardScaler
        scaler = StandardScaler()
        self.df_normalized = self.df.copy()
        self.df_normalized[numerical_cols] = scaler.fit_transform(self.df[numerical_cols])
        self.scaler = scaler
        
        print(f"\n✓ Đã chuẩn hóa {len(numerical_cols)} biến số bằng StandardScaler")
        
        # 4. Thống kê SAU chuẩn hóa
        print("\n" + "="*80)
        print("SAU CHUẨN HÓA / AFTER NORMALIZATION")
        print("="*80)
        stats_after = self.df_normalized[numerical_cols].describe()
        print("\nThống kê mô tả:")
        print(stats_after)
        
        self.report.append("\n" + "="*80)
        self.report.append("SAU CHUẨN HÓA / AFTER NORMALIZATION")
        self.report.append("="*80)
        self.report.append("\nThống kê mô tả:")
        self.report.append(str(stats_after))
        
        # 5. So sánh chi tiết từng biến
        print("\n" + "="*80)
        print("SO SÁNH CHI TIẾT / DETAILED COMPARISON")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("SO SÁNH CHI TIẾT / DETAILED COMPARISON")
        self.report.append("="*80)
        
        comparison_data = []
        for col in numerical_cols:
            before_mean = self.df[col].mean()
            before_std = self.df[col].std()
            before_min = self.df[col].min()
            before_max = self.df[col].max()
            
            after_mean = self.df_normalized[col].mean()
            after_std = self.df_normalized[col].std()
            after_min = self.df_normalized[col].min()
            after_max = self.df_normalized[col].max()
            
            print(f"\n{col}:")
            print(f"  Trước / Before: Mean={before_mean:.4f}, Std={before_std:.4f}, Min={before_min:.4f}, Max={before_max:.4f}")
            print(f"  Sau / After:   Mean={after_mean:.4f}, Std={after_std:.4f}, Min={after_min:.4f}, Max={after_max:.4f}")
            
            self.report.append(f"\n{col}:")
            self.report.append(f"  Trước / Before: Mean={before_mean:.4f}, Std={before_std:.4f}, Min={before_min:.4f}, Max={before_max:.4f}")
            self.report.append(f"  Sau / After:   Mean={after_mean:.4f}, Std={after_std:.4f}, Min={after_min:.4f}, Max={after_max:.4f}")
        
        # 6. Thông tin mã hóa biến phân loại
        print(f"\n✓ Thông tin mã hóa biến phân loại / Categorical encoding:")
        self.report.append(f"\n✓ Thông tin mã hóa biến phân loại / Categorical encoding:")
        
        for col, le in self.le_dict.items():
            print(f"  {col}: {len(le.classes_)} lớp")
            print(f"    Ánh xạ / Mapping: {dict(zip(le.classes_, le.transform(le.classes_)))}")
            self.report.append(f"  {col}: {len(le.classes_)} lớp")
            self.report.append(f"    Ánh xạ / Mapping: {dict(zip(le.classes_, le.transform(le.classes_)))}")
        
        self.analysis_history.append("Chuẩn Hóa Dữ Liệu / Data Normalization")
    
    # ============ A) DESCRIPTIVE STATISTICS ============
    def descriptive_statistics(self):
        """
        TÍNH NĂNG / FEATURE: Phân Tích Thống Kê Mô Tả / Descriptive Statistics Analysis
        MÔ TẢ / DESCRIPTION: Cung cấp tóm tắt toàn diện về tập dữ liệu bao gồm / Provides comprehensive dataset summary including:
        - Thông tin cơ bản (số dòng, số cột, tên cột) / Basic information (rows, columns, names)
        - Kiểu dữ liệu của tất cả biến / Data types of all variables
        - Phát hiện giá trị thiếu / Missing values detection
        - Tóm tắt thống kê cho các biến định lượng / Statistical summary for numerical variables
        - Đếm giá trị cho các biến định tính / Value counts for categorical variables
        """
        print("\n" + "="*80)
        print("A) PHÂN TÍCH THỐNG KÊ MÔ TẢ / DESCRIPTIVE STATISTICS ANALYSIS")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("A) PHÂN TÍCH THỐNG KÊ MÔ TẢ / DESCRIPTIVE STATISTICS ANALYSIS")
        self.report.append("="*80)
        
        # 1. Basic information
        print("\n1. THÔNG TIN CƠ BẢN / BASIC DATA INFORMATION")
        self.report.append("\n1. THÔNG TIN CƠ BẢN / BASIC DATA INFORMATION")
        print(f"   - Số dòng / Number of rows: {self.df.shape[0]}")
        print(f"   - Số cột / Number of columns: {self.df.shape[1]}")
        print(f"   - Tên cột / Column names: {list(self.df.columns)}")
        self.report.append(f"   - Số dòng / Number of rows: {self.df.shape[0]}")
        self.report.append(f"   - Số cột / Number of columns: {self.df.shape[1]}")
        
        # 2. Data types
        print("\n2. KIỂU DỮ LIỆU / DATA TYPES")
        self.report.append("\n2. KIỂU DỮ LIỆU / DATA TYPES")
        print(self.df.dtypes)
        self.report.append(str(self.df.dtypes))
        
        # 3. Missing values
        print("\n3. KIỂM TRA GIÁ TRỊ THIẾU / MISSING VALUES CHECK")
        self.report.append("\n3. KIỂM TRA GIÁ TRỊ THIẾU / MISSING VALUES CHECK")
        missing = self.df.isnull().sum()
        if missing.sum() == 0:
            print("   - Không phát hiện giá trị thiếu / No missing values detected")
            self.report.append("   - Không phát hiện giá trị thiếu / No missing values detected")
        else:
            print(missing[missing > 0])
            self.report.append(str(missing[missing > 0]))
        
        # 4. Numerical statistics
        print("\n4. THỐNG KÊ BIẾN ĐỊNH LƯỢNG / NUMERICAL VARIABLES STATISTICS")
        self.report.append("\n4. THỐNG KÊ BIẾN ĐỊNH LƯỢNG / NUMERICAL VARIABLES STATISTICS")
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns
        stats_df = self.df[numerical_cols].describe()
        print(stats_df)
        self.report.append("\n" + str(stats_df))
        
        # 5. Categorical statistics
        print("\n5. THỐNG KÊ BIẾN ĐỊNH TÍNH / CATEGORICAL VARIABLES STATISTICS")
        self.report.append("\n5. THỐNG KÊ BIẾN ĐỊNH TÍNH / CATEGORICAL VARIABLES STATISTICS")
        categorical_cols = self.df.select_dtypes(include=['object']).columns
        for col in categorical_cols:
            print(f"\n   {col}:")
            print(self.df[col].value_counts())
            self.report.append(f"\n   {col}:")
            self.report.append(str(self.df[col].value_counts()))
        
        self.analysis_history.append("Phân Tích Thống Kê Mô Tả / Descriptive Statistics")
    
    # ============ B) ESTIMATION & HYPOTHESIS TESTING ============
    def estimation_hypothesis_testing(self):
        """
        TÍNH NĂNG / FEATURE: Ước Lượng và Kiểm Định / Estimation and Hypothesis Testing
        MÔ TẢ / DESCRIPTION: Thực hiện suy diễn thống kê bao gồm / Performs statistical inference including:
        - Khoảng tin cậy 95% / 95% Confidence intervals
        - Kiểm định T / T-tests
        - Kiểm định Chi-bình phương / Chi-square tests
        - Giải thích giá trị P / P-value interpretation
        """
        print("\n" + "="*80)
        print("B) ƯỚC LƯỢNG VÀ KIỂM ĐỊNH / ESTIMATION AND HYPOTHESIS TESTING")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("B) ƯỚC LƯỢNG VÀ KIỂM ĐỊNH / ESTIMATION AND HYPOTHESIS TESTING")
        self.report.append("="*80)
        
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns
        
        # 1. Confidence intervals for means
        print("\n1. KHOẢNG TIN CẬY 95% CHO TRUNG BÌNH / 95% CONFIDENCE INTERVALS FOR MEANS")
        self.report.append("\n1. KHOẢNG TIN CẬY 95% CHO TRUNG BÌNH / 95% CONFIDENCE INTERVALS FOR MEANS")
        for col in numerical_cols:
            data = self.df[col].dropna()
            mean = data.mean()
            std = data.std()
            n = len(data)
            se = std / np.sqrt(n)
            ci = stats.t.interval(0.95, n-1, loc=mean, scale=se)
            print(f"\n   {col}:")
            print(f"      Mean: {mean:.4f}")
            print(f"      95% CI: [{ci[0]:.4f}, {ci[1]:.4f}]")
            self.report.append(f"\n   {col}: Mean={mean:.4f}, CI=[{ci[0]:.4f}, {ci[1]:.4f}]")
        
        # 2. T-test: age by deposit status
        print("\n2. KIỂM ĐỊNH T: TUỔI THEO TRẠNG THÁI GỬI TIỀN / T-TEST: AGE BY DEPOSIT STATUS")
        self.report.append("\n2. KIỂM ĐỊNH T: TUỔI THEO TRẠNG THÁI GỬI TIỀN / T-TEST: AGE BY DEPOSIT STATUS")
        if 'deposit' in self.df.columns and 'age' in self.df.columns:
            yes_group = self.df[self.df['deposit'] == 'yes']['age']
            no_group = self.df[self.df['deposit'] == 'no']['age']
            t_stat, p_value = stats.ttest_ind(yes_group, no_group)
            print(f"   Nhóm Có (n={len(yes_group)}): Trung bình={yes_group.mean():.4f}, Độ lệch chuẩn={yes_group.std():.4f}")
            print(f"   Nhóm Không (n={len(no_group)}): Trung bình={no_group.mean():.4f}, Độ lệch chuẩn={no_group.std():.4f}")
            print(f"   T-statistic: {t_stat:.4f}")
            print(f"   P-value: {p_value:.6f}")
            sig = "Có ý nghĩa thống kê / Statistically significant" if p_value < 0.05 else "Không có ý nghĩa / Not statistically significant"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   T-stat={t_stat:.4f}, P-value={p_value:.6f}")
        
        # 3. Chi-square test
        print("\n3. KIỂM ĐỊNH CHI-BÌNH PHƯƠNG: CÔNG VIỆC vs GỬI TIỀN / CHI-SQUARE TEST: JOB vs DEPOSIT")
        self.report.append("\n3. KIỂM ĐỊNH CHI-BÌNH PHƯƠNG: CÔNG VIỆC vs GỬI TIỀN / CHI-SQUARE TEST: JOB vs DEPOSIT")
        if 'job' in self.df.columns and 'deposit' in self.df.columns:
            ct = pd.crosstab(self.df['job'], self.df['deposit'])
            chi2, p_value, dof, expected = stats.chi2_contingency(ct)
            print(f"   Chi-square: {chi2:.4f}")
            print(f"   P-value: {p_value:.6f}")
            print(f"   Degrees of freedom: {dof}")
            sig = "Có sự liên kết / Association exists" if p_value < 0.05 else "Không có sự liên kết / No association"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   Chi-square={chi2:.4f}, P-value={p_value:.6f}")
        
        self.analysis_history.append("Ước Lượng & Kiểm Định / Estimation & Hypothesis Testing")
    
    # ============ C) CORRELATION ANALYSIS ============
    def correlation_analysis(self):
        """
        TÍNH NĂNG / FEATURE: Phân Tích Tương Quan / Correlation Analysis
        MÔ TẢ / DESCRIPTION: Kiểm tra các mối quan hệ / Examines relationships including:
        - Ma trận tương quan Pearson / Pearson correlation matrix
        - Tương quan với biến mục tiêu / Correlation with target variable
        - Tương quan hạng Spearman / Spearman rank correlation
        """
        print("\n" + "="*80)
        print("C) PHÂN TÍCH TƯƠNG QUAN / CORRELATION ANALYSIS")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("C) PHÂN TÍCH TƯƠNG QUAN / CORRELATION ANALYSIS")
        self.report.append("="*80)
        
        numerical_cols = self.df_encoded.select_dtypes(include=[np.number]).columns
        
        # 1. Pearson correlation
        print("\n1. MA TRẬN TƯƠNG QUAN PEARSON / PEARSON CORRELATION MATRIX")
        self.report.append("\n1. MA TRẬN TƯƠNG QUAN PEARSON / PEARSON CORRELATION MATRIX")
        corr_matrix = self.df_encoded[numerical_cols].corr()
        print(corr_matrix)
        self.report.append("\n" + str(corr_matrix))
        
        # 2. Correlation with target variable (deposit)
        if 'deposit' in self.df_encoded.columns:
            print("\n2. TƯƠNG QUAN VỚI BIẾN MỤC TIÊU (GỬI TIỀN) / CORRELATION WITH TARGET VARIABLE (DEPOSIT)")
            self.report.append("\n2. TƯƠNG QUAN VỚI BIẾN MỤC TIÊU (GỬI TIỀN) / CORRELATION WITH TARGET VARIABLE (DEPOSIT)")
            deposit_corr = corr_matrix['deposit'].sort_values(ascending=False)
            print(deposit_corr)
            self.report.append("\n" + str(deposit_corr))
        
        # 3. Spearman correlation for age vs balance
        print("\n3. TƯƠNG QUAN HẠNG SPEARMAN: TUỔI vs SỐ DƯ / SPEARMAN CORRELATION: AGE vs BALANCE")
        self.report.append("\n3. TƯƠNG QUAN HẠNG SPEARMAN: TUỔI vs SỐ DƯ / SPEARMAN CORRELATION: AGE vs BALANCE")
        if 'age' in self.df.columns and 'balance' in self.df.columns:
            spearman_corr, p_value = stats.spearmanr(self.df['age'], self.df['balance'])
            print(f"   Hệ số tương quan: {spearman_corr:.4f}")
            print(f"   P-value: {p_value:.6f}")
            self.report.append(f"   Hệ số tương quan: {spearman_corr:.4f}, P-value={p_value:.6f}")
        
        self.analysis_history.append("Phân Tích Tương Quan / Correlation Analysis")
    
    # ============ D) ANOVA ANALYSIS ============
    def anova_analysis(self):
        """
        TÍNH NĂNG / FEATURE: Phân Tích ANOVA / ANOVA Analysis
        MÔ TẢ / DESCRIPTION: Kiểm tra sự khác biệt giữa / Tests differences between:
        - One-way ANOVA: Tuổi theo công việc / Age by job type
        - One-way ANOVA: Số dư theo hôn nhân / Balance by marital status
        - F-statistic và P-value / F-statistic and P-value
        """
        print("\n" + "="*80)
        print("D) PHÂN TÍCH ANOVA / ANOVA (ANALYSIS OF VARIANCE)")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("D) PHÂN TÍCH ANOVA / ANOVA (ANALYSIS OF VARIANCE)")
        self.report.append("="*80)
        
        # 1. One-way ANOVA: Age by job
        print("\n1. ANOVA MỘT CHIỀU: TUỔI THEO LOẠI CÔNG VIỆC / ONE-WAY ANOVA: AGE BY JOB TYPE")
        self.report.append("\n1. ANOVA MỘT CHIỀU: TUỔI THEO LOẠI CÔNG VIỆC / ONE-WAY ANOVA: AGE BY JOB TYPE")
        if 'age' in self.df.columns and 'job' in self.df.columns:
            groups = self.df.groupby('job')['age'].apply(list)
            f_stat, p_value = stats.f_oneway(*groups)
            print(f"   F-statistic: {f_stat:.4f}")
            print(f"   P-value: {p_value:.6f}")
            sig = "Có sự khác biệt đáng kể / Significant differences exist" if p_value < 0.05 else "Không có sự khác biệt / No significant differences"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   F-stat={f_stat:.4f}, P-value={p_value:.6f}")
            
            print("\n   Trung bình tuổi theo loại công việc:")
            self.report.append("\n   Trung bình tuổi theo loại công việc:")
            job_means = self.df.groupby('job')['age'].mean().sort_values(ascending=False)
            print(job_means)
            self.report.append("\n" + str(job_means))
        
        # 2. One-way ANOVA: Balance by marital status
        print("\n2. ANOVA MỘT CHIỀU: SỐ DƯ THEO TÌNH TRẠNG HÔN NHÂN / ONE-WAY ANOVA: BALANCE BY MARITAL STATUS")
        self.report.append("\n2. ANOVA MỘT CHIỀU: SỐ DƯ THEO TÌNH TRẠNG HÔN NHÂN / ONE-WAY ANOVA: BALANCE BY MARITAL STATUS")
        if 'balance' in self.df.columns and 'marital' in self.df.columns:
            groups = self.df.groupby('marital')['balance'].apply(list)
            f_stat, p_value = stats.f_oneway(*groups)
            print(f"   F-statistic: {f_stat:.4f}")
            print(f"   P-value: {p_value:.6f}")
            self.report.append(f"   F-stat={f_stat:.4f}, P-value={p_value:.6f}")
        
        self.analysis_history.append("Phân Tích ANOVA / ANOVA Analysis")
    
    # ============ E) REGRESSION ANALYSIS ============
    def regression_analysis(self):
        """
        TÍNH NĂNG / FEATURE: Phân Tích Hồi Quy / Regression Analysis
        MÔ TẢ / DESCRIPTION: Mô hình hóa các mối quan hệ / Models relationships including:
        - Hồi quy tuyến tính đơn / Simple Linear Regression
        - Hồi quy tuyến tính đa / Multiple Linear Regression
        - R-squared / R-squared values
        """
        print("\n" + "="*80)
        print("E) PHÂN TÍCH HỒI QUY / REGRESSION ANALYSIS")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("E) PHÂN TÍCH HỒI QUY / REGRESSION ANALYSIS")
        self.report.append("="*80)
        
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numerical_cols) >= 2:
            # 1. Simple linear regression: balance vs age
            print("\n1. HỒI QUY TUYẾN TÍNH ĐƠN: SỐ DƯ vs TUỔI / SIMPLE LINEAR REGRESSION: BALANCE vs AGE")
            self.report.append("\n1. HỒI QUY TUYẾN TÍNH ĐƠN: SỐ DƯ vs TUỔI / SIMPLE LINEAR REGRESSION: BALANCE vs AGE")
            if 'age' in self.df.columns and 'balance' in self.df.columns:
                X_simple = self.df[['age']].values
                y_simple = self.df['balance'].values
                model_simple = LinearRegression()
                model_simple.fit(X_simple, y_simple)
                r2_simple = model_simple.score(X_simple, y_simple)
                
                print(f"   Hệ số chặn: {model_simple.intercept_:.4f}")
                print(f"   Hệ số (tuổi): {model_simple.coef_[0]:.4f}")
                print(f"   R-bình phương: {r2_simple:.4f}")
                print(f"   Phương trình: Số dư = {model_simple.intercept_:.4f} + {model_simple.coef_[0]:.4f} * Tuổi")
                self.report.append(f"   Hệ số chặn={model_simple.intercept_:.4f}")
                self.report.append(f"   Hệ số={model_simple.coef_[0]:.4f}")
                self.report.append(f"   R-bình phương={r2_simple:.4f}")
            
            # 2. Multiple linear regression
            print("\n2. HỒI QUY TUYẾN TÍNH ĐA / MULTIPLE LINEAR REGRESSION")
            self.report.append("\n2. HỒI QUY TUYẾN TÍNH ĐA / MULTIPLE LINEAR REGRESSION")
            selected_cols = [col for col in numerical_cols if col != 'balance'][:5]
            if 'balance' in numerical_cols and len(selected_cols) > 0:
                X_multi = self.df[selected_cols].values
                y_multi = self.df['balance'].values
                model_multi = LinearRegression()
                model_multi.fit(X_multi, y_multi)
                r2_multi = model_multi.score(X_multi, y_multi)
                
                print(f"   Biến độc lập: {selected_cols}")
                print(f"   Hệ số chặn: {model_multi.intercept_:.4f}")
                print(f"   Các hệ số:")
                for i, col in enumerate(selected_cols):
                    print(f"      {col}: {model_multi.coef_[i]:.4f}")
                print(f"   R-bình phương: {r2_multi:.4f}")
                self.report.append(f"   Biến độc lập: {selected_cols}")
                self.report.append(f"   Hệ số chặn: {model_multi.intercept_:.4f}")
                self.report.append(f"   R-bình phương: {r2_multi:.4f}")
        
        self.analysis_history.append("Phân Tích Hồi Quy / Regression Analysis")
    
    def run_all_analysis(self):
        """Run all analyses - Normalization MUST run first"""
        print("\n" + "="*80)
        print("CHẠY TẤT CẢ PHÂN TÍCH / RUNNING ALL ANALYSES")
        print("="*80)
        print("\n⚠ Lưu ý / NOTE: Chuẩn hóa dữ liệu sẽ chạy TRƯỚC / Normalization will run FIRST")
        
        # Normalize data first
        self.normalize_data()
        
        # Then run all analyses
        self.descriptive_statistics()
        self.estimation_hypothesis_testing()
        self.correlation_analysis()
        self.anova_analysis()
        self.regression_analysis()

    def display_analysis_summary(self):
        """Display summary of completed analyses"""
        print("\n" + "="*80)
        print("TÓM TẮT PHÂN TÍCH ĐÃ HOÀN THÀNH / ANALYSIS SUMMARY")
        print("="*80)
        print(f"\nTổng số phân tích đã chạy / Total analyses completed: {len(self.analysis_history)}")
        print("\nCác phân tích đã hoàn thành / Completed analyses:")
        for i, analysis in enumerate(self.analysis_history, 1):
            print(f"   {i}. {analysis}")
        print("\n" + "="*80)
    
    def save_report(self, filename='bank_analysis_report.txt', file_format='txt'):
        """Save report to file in specified format (txt, docx, pdf)"""
        file_format = file_format.lower()
        
        if file_format == 'txt':
            self._save_txt(filename)
        elif file_format == 'docx':
            self._save_docx(filename)
        elif file_format == 'pdf':
            self._save_pdf(filename)
        else:
            print(f"Unsupported format: {file_format}")
            print("Supported formats: txt, docx, pdf")
    
    def _save_txt(self, filename):
        """Save report as TXT file"""
        if not filename.endswith('.txt'):
            filename = filename.replace('.docx', '').replace('.pdf', '') + '.txt'
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.report))
            print(f"\n✓ TXT report saved to: {filename}")
        except Exception as e:
            print(f"✗ Error saving TXT file: {e}")
    
    def _save_docx(self, filename):
        """Save report as DOCX file"""
        if not DOCX_AVAILABLE:
            print("✗ python-docx not installed. Run: pip install python-docx")
            return
        
        if not filename.endswith('.docx'):
            filename = filename.replace('.txt', '').replace('.pdf', '') + '.docx'
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('BANK MARKETING ANALYSIS REPORT', 0)
            title.alignment = 1
            
            # Add timestamp
            timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            timestamp.alignment = 1
            
            doc.add_paragraph()
            
            # Add content
            for line in self.report:
                if line.startswith('='):
                    heading = doc.add_heading(line.strip('=').strip(), level=1)
                    heading.alignment = 1
                elif any(line.startswith(x) for x in ['A)', 'B)', 'C)', 'D)', 'E)']):
                    section = doc.add_heading(line, level=2)
                    section.runs[0].font.size = Pt(12)
                elif line.startswith('   '):
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc.save(filename)
            print(f"\n✓ DOCX report saved to: {filename}")
        except Exception as e:
            print(f"✗ Error saving DOCX file: {e}")
    
    def _save_pdf(self, filename):
        """Save report as PDF file"""
        if not PDF_AVAILABLE:
            print("✗ reportlab not installed. Run: pip install reportlab")
            return
        
        if not filename.endswith('.pdf'):
            filename = filename.replace('.txt', '').replace('.docx', '') + '.pdf'
        
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            
            pdf = SimpleDocTemplate(filename, pagesize=letter,
                                  rightMargin=0.5*inch, leftMargin=0.5*inch,
                                  topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            story = []
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=12,
                spaceBefore=12,
                alignment=TA_LEFT
            )
            
            title = Paragraph("BANK MARKETING ANALYSIS REPORT", title_style)
            story.append(title)
            
            timestamp_text = f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
            timestamp = Paragraph(timestamp_text, styles['Normal'])
            story.append(timestamp)
            story.append(Spacer(1, 0.3*inch))
            
            for line in self.report:
                if line.strip() == '':
                    story.append(Spacer(1, 0.1*inch))
                elif line.startswith('='):
                    heading = Paragraph(line.strip('=').strip(), heading_style)
                    story.append(heading)
                else:
                    line_safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p = Paragraph(line_safe, styles['Normal'])
                    story.append(p)
            
            pdf.build(story)
            print(f"\n✓ PDF report saved to: {filename}")
        except Exception as e:
            print(f"✗ Error saving PDF file: {e}")
    
    def _export_menu(self):
        """Display export format menu"""
        self.display_analysis_summary()
        
        print("\n" + "="*80)
        print("XUẤT KẾT QUẢ PHÂN TÍCH / EXPORT ANALYSIS RESULTS")
        print("="*80)
        print("Chọn định dạng xuất file / Select export format:")
        print("1. TXT (Tệp văn bản / Text file)")
        print("2. DOCX (Tài liệu Word / Word document)")
        print("3. PDF (Định dạng PDF / Portable Document Format)")
        print("4. Tất cả định dạng / All formats (Xuất tất cả 3 / export to all three)")
        print("0. Quay lại / Back to main menu")
        
        choice = input("\nNhập lựa chọn / Enter choice (0-4): ").strip()
        
        base_filename = 'bank_analysis_report'
        
        if choice == '1':
            self.save_report(f"{base_filename}.txt", 'txt')
        elif choice == '2':
            self.save_report(f"{base_filename}.docx", 'docx')
        elif choice == '3':
            self.save_report(f"{base_filename}.pdf", 'pdf')
        elif choice == '4':
            print("\nĐang xuất sang tất cả các định dạng / Exporting to all formats...")
            self.save_report(f"{base_filename}.txt", 'txt')
            self.save_report(f"{base_filename}.docx", 'docx')
            self.save_report(f"{base_filename}.pdf", 'pdf')
            print("\n✓ Tất cả báo cáo đã được xuất thành công / All reports exported successfully!")
        elif choice == '0':
            pass
        else:
            print("\nLựa chọn không hợp lệ / Invalid choice!")

    def display_menu(self):
        """Display interactive menu"""
        # Check if normalization has been done
        is_normalized = hasattr(self, 'df_normalized')
        
        while True:
            print("\n" + "="*80)
            print("HỆ THỐNG PHÂN TÍCH THỐNG KÊ MARKETING NGÂN HÀNG / BANK MARKETING STATISTICAL ANALYSIS SYSTEM")
            print("="*80)
            
            # Show normalization status
            norm_status = "✓ ĐÃ CHUẨN HÓA" if is_normalized else "⚠ CHƯA CHUẨN HÓA"
            print(f"\nTrạng thái chuẩn hóa / Normalization status: {norm_status}")
            
            print("\nChọn phân tích cần thực hiện / Select analysis to perform:")
            print("0. Chuẩn Hóa Dữ Liệu / Data Normalization - Chuẩn hóa các biến số / Normalize numerical variables")
            print("1. Thống Kê Mô Tả / Descriptive Statistics - Thông tin cơ bản về dữ liệu / Basic data information")
            print("2. Ước Lượng & Kiểm Định / Estimation & Hypothesis Testing - Khoảng tin cậy và kiểm định / Confidence intervals and tests")
            print("3. Phân Tích Tương Quan / Correlation Analysis - Mối quan hệ giữa các biến / Relationships between variables")
            print("4. Phân Tích ANOVA / ANOVA Analysis - So sánh trung bình nhóm / Compare group means")
            print("5. Phân Tích Hồi Quy / Regression Analysis - Mô hình dự đoán / Predictive models")
            print("6. Chạy Tất Cả Phân Tích / Run All Analyses - Thực hiện toàn bộ phân tích / Run all analyses (includes normalization)")
            print("7. Xuất Kết Quả / Export Results - Lưu báo cáo / Save reports")
            print("8. Thoát / Exit - Kết thúc chương trình / End program")
            
            choice = input("\nNhập lựa chọn / Enter choice (0-8): ").strip()
            
            if choice == '0':
                self.normalize_data()
                is_normalized = True
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '1':
                if not is_normalized:
                    print("\n⚠ LỖI / ERROR: Bạn phải chuẩn hóa dữ liệu trước! / You must normalize data first!")
                    print("   Vui lòng chọn mục 0 / Please select option 0")
                    input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
                    continue
                self.descriptive_statistics()
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '2':
                if not is_normalized:
                    print("\n⚠ LỖI / ERROR: Bạn phải chuẩn hóa dữ liệu trước! / You must normalize data first!")
                    print("   Vui lòng chọn mục 0 / Please select option 0")
                    input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
                    continue
                self.estimation_hypothesis_testing()
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '3':
                if not is_normalized:
                    print("\n⚠ LỖI / ERROR: Bạn phải chuẩn hóa dữ liệu trước! / You must normalize data first!")
                    print("   Vui lòng chọn mục 0 / Please select option 0")
                    input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
                    continue
                self.correlation_analysis()
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '4':
                if not is_normalized:
                    print("\n⚠ LỖI / ERROR: Bạn phải chuẩn hóa dữ liệu trước! / You must normalize data first!")
                    print("   Vui lòng chọn mục 0 / Please select option 0")
                    input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
                    continue
                self.anova_analysis()
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '5':
                if not is_normalized:
                    print("\n⚠ LỖI / ERROR: Bạn phải chuẩn hóa dữ liệu trước! / You must normalize data first!")
                    print("   Vui lòng chọn mục 0 / Please select option 0")
                    input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
                    continue
                self.regression_analysis()
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '6':
                print("\nChạy tất cả các phân tích / Running all analyses...")
                self.run_all_analysis()
                is_normalized = True
                input("\nNhấn Enter để tiếp tục / Press Enter to continue...")
            elif choice == '7':
                self._export_menu()
            elif choice == '8':
                print("\nCảm ơn đã sử dụng hệ thống phân tích / Thank you for using the analysis system!")
                break
            else:
                print("\nLựa chọn không hợp lệ / Invalid choice. Vui lòng thử lại / Please try again!")

def main():
    """Main function"""
    filepath = 'bank.csv'
    
    try:
        analysis = BankMarketingAnalysis(filepath)
        analysis.display_menu()
    except FileNotFoundError:
        print(f"Error: File not found: {filepath}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()

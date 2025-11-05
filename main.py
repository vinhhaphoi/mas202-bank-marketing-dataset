import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import LabelEncoder, StandardScaler
import warnings
from datetime import datetime
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

warnings.filterwarnings('ignore')

class BankMarketingAnalysis:
    def __init__(self, filepath):
        """Khởi tạo phân tích với dữ liệu từ file CSV"""
        self.df = pd.read_csv(filepath)
        self.report = []
        self.analysis_history = []
        self.insights = []
        self.visualizations = []
        
        # Tạo thư mục output
        self.output_dir = 'analysis_output'
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # Cấu hình pandas để hiển thị đầy đủ
        pd.set_option('display.max_columns', None)
        pd.set_option('display.width', None)
        pd.set_option('display.max_colwidth', None)
        pd.set_option('display.expand_frame_repr', False)
        
        # Cấu hình matplotlib
        plt.style.use('seaborn-v0_8-darkgrid')
        sns.set_palette("husl")
        
        # Khởi tạo mô tả biến
        self._init_variable_descriptions()
        
        self.setup_data()

    def _init_variable_descriptions(self):
        """Khởi tạo mô tả các biến để dùng trong báo cáo và chú thích biểu đồ"""
        # Các mô tả mẫu cho các biến thường gặp (bổ sung khi cần)
        self.var_desc = {
            'age': 'Tuổi của khách hàng',
            'balance': 'Số dư tài khoản (đơn vị tiền tệ)',
            'day': 'Ngày trong tháng khi liên hệ',
            'duration': 'Thời lượng cuộc gọi (giây)',
            'campaign': 'Số lần liên hệ trong chiến dịch hiện tại',
            'pdays': 'Số ngày kể từ lần liên hệ trước (-1 nếu chưa từng liên hệ)',
            'previous': 'Số lần liên hệ trước đó',
            'job': 'Nghề nghiệp của khách hàng',
            'marital': 'Tình trạng hôn nhân',
            'deposit': 'Khách hàng gửi tiền tiết kiệm hay không (yes/no)'
        }
        # Thêm mô tả cho mọi cột còn lại dưới dạng chung
        for col in self.df.columns:
            if col not in self.var_desc:
                self.var_desc[col] = f'Mô tả chưa có cho biến "{col}"'

    def save_variable_descriptions(self):
        """Lưu file mô tả biến (variable_descriptions.txt) vào thư mục output"""
        filename = f'{self.output_dir}/variable_descriptions.txt'
        lines = ["MÔ TẢ CÁC BIẾN (Tiếng Việt):", "="*60]
        for col, desc in self.var_desc.items():
            lines.append(f"{col}: {desc}")
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        self.report.append(f"\n✓ Mô tả biến được lưu vào: {filename}")
        print(f"\n✓ Đã lưu mô tả biến: {filename}")
        return filename

    def setup_data(self):
        """Chuẩn bị dữ liệu cho phân tích"""
        self.le_dict = {}
        self.df_encoded = self.df.copy()
        
        categorical_cols = self.df_encoded.select_dtypes(include=['object']).columns
        for col in categorical_cols:
            le = LabelEncoder()
            self.df_encoded[col] = le.fit_transform(self.df_encoded[col])
            self.le_dict[col] = le
    
    def normalize_data(self):
        """
        TÍNH NĂNG: Chuẩn Hóa Dữ Liệu
        MÔ TẢ: Chuẩn hóa các biến định lượng bằng StandardScaler
        - Chuyển đổi tất cả biến số về thang đo [trung bình=0, độ lệch chuẩn=1]
        - Giúp các thuật toán học máy hoạt động tốt hơn
        - Lưu bộ chuẩn hóa để chuyển đổi dữ liệu mới
        """
        print("\n" + "="*80)
        print("CHUẨN HÓA DỮ LIỆU")
        print("="*80)
        
        # 1. Lấy các biến số
        numerical_cols = list(self.df.select_dtypes(include=[np.number]).columns)
        
        print(f"\n✓ Phát hiện {len(numerical_cols)} biến số cần chuẩn hóa:")
        print(f"  {numerical_cols}")
        
        # 2. Thống kê TRƯỚC chuẩn hóa
        print("\n" + "="*80)
        print("TRƯỚC CHUẨN HÓA")
        print("="*80)
        stats_before = self.df[numerical_cols].describe()
        print("\nThống kê mô tả:")
        print(stats_before)
        
        self.report.append("\n" + "="*80)
        self.report.append("CHUẨN HÓA DỮ LIỆU")
        self.report.append("="*80)
        self.report.append("\n" + "="*80)
        self.report.append("TRƯỚC CHUẨN HÓA")
        self.report.append("="*80)
        self.report.append("\nThống kê mô tả:")
        self.report.append(str(stats_before))
        
        # 3. Áp dụng StandardScaler
        scaler = StandardScaler()
        self.df_normalized = self.df.copy()
        self.df_normalized[numerical_cols] = scaler.fit_transform(self.df[numerical_cols])
        self.scaler = scaler
        
        print(f"\n✓ Đã chuẩn hóa {len(numerical_cols)} biến số bằng StandardScaler")
        
        # 4. Thống kê SAU chuẩn hóa
        print("\n" + "="*80)
        print("SAU CHUẨN HÓA")
        print("="*80)
        stats_after = self.df_normalized[numerical_cols].describe()
        print("\nThống kê mô tả:")
        print(stats_after)
        
        self.report.append("\n" + "="*80)
        self.report.append("SAU CHUẨN HÓA")
        self.report.append("="*80)
        self.report.append("\nThống kê mô tả:")
        self.report.append(str(stats_after))
        
        # 5. So sánh chi tiết từng biến
        print("\n" + "="*80)
        print("SO SÁNH CHI TIẾT")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("SO SÁNH CHI TIẾT")
        self.report.append("="*80)
        
        comparison_data = []
        for col in numerical_cols:
            before_mean = self.df[col].mean()
            before_std = self.df[col].std()
            before_min = self.df[col].min()
            before_max = self.df[col].max()
            
            after_mean = self.df_normalized[col].mean()
            after_std = self.df_normalized[col].std()
            after_min = self.df_normalized[col].min()
            after_max = self.df_normalized[col].max()
            
            print(f"\n{col}:")
            print(f"  Trước: Trung bình={before_mean:.4f}, Độ lệch chuẩn={before_std:.4f}, Min={before_min:.4f}, Max={before_max:.4f}")
            print(f"  Sau:   Trung bình={after_mean:.4f}, Độ lệch chuẩn={after_std:.4f}, Min={after_min:.4f}, Max={after_max:.4f}")
            
            self.report.append(f"\n{col}:")
            self.report.append(f"  Trước: Trung bình={before_mean:.4f}, Độ lệch chuẩn={before_std:.4f}, Min={before_min:.4f}, Max={before_max:.4f}")
            self.report.append(f"  Sau:   Trung bình={after_mean:.4f}, Độ lệch chuẩn={after_std:.4f}, Min={after_min:.4f}, Max={after_max:.4f}")
        
        # 6. Thông tin mã hóa biến phân loại
        print(f"\n✓ Thông tin mã hóa biến phân loại:")
        self.report.append(f"\n✓ Thông tin mã hóa biến phân loại:")
        
        for col, le in self.le_dict.items():
            print(f"  {col}: {len(le.classes_)} lớp")
            print(f"    Ánh xạ: {dict(zip(le.classes_, le.transform(le.classes_)))}")
            self.report.append(f"  {col}: {len(le.classes_)} lớp")
            self.report.append(f"    Ánh xạ: {dict(zip(le.classes_, le.transform(le.classes_)))}")
        
        self.analysis_history.append("Chuẩn Hóa Dữ Liệu")
    
    # ============ A) DESCRIPTIVE STATISTICS ============
    def descriptive_statistics(self):
        """
        TÍNH NĂNG: Phân Tích Thống Kê Mô Tả
        MÔ TẢ: Cung cấp tóm tắt toàn diện về tập dữ liệu bao gồm:
        - Thông tin cơ bản (số dòng, số cột, tên cột)
        - Kiểu dữ liệu của tất cả biến
        - Phát hiện giá trị thiếu
        - Tóm tắt thống kê cho các biến định lượng
        - Đếm giá trị cho các biến định tính
        """
        print("\n" + "="*80)
        print("A) PHÂN TÍCH THỐNG KÊ MÔ TẢ")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("A) PHÂN TÍCH THỐNG KÊ MÔ TẢ")
        self.report.append("="*80)
        
        # 1. Basic information
        print("\n1. THÔNG TIN CƠ BẢN")
        self.report.append("\n1. THÔNG TIN CƠ BẢN")
        print(f"   - Số dòng: {self.df.shape[0]}")
        print(f"   - Số cột: {self.df.shape[1]}")
        print(f"   - Tên cột: {list(self.df.columns)}")
        self.report.append(f"   - Số dòng: {self.df.shape[0]}")
        self.report.append(f"   - Số cột: {self.df.shape[1]}")
        
        # 2. Data types
        print("\n2. KIỂU DỮ LIỆU")
        self.report.append("\n2. KIỂU DỮ LIỆU")
        print(self.df.dtypes)
        self.report.append(str(self.df.dtypes))
        
        # 3. Missing values
        print("\n3. KIỂM TRA GIÁ TRỊ THIẾU")
        self.report.append("\n3. KIỂM TRA GIÁ TRỊ THIẾU")
        missing = self.df.isnull().sum()
        if missing.sum() == 0:
            print("   - Không phát hiện giá trị thiếu")
            self.report.append("   - Không phát hiện giá trị thiếu")
        else:
            print(missing[missing > 0])
            self.report.append(str(missing[missing > 0]))
        
        # 4. Numerical statistics
        print("\n4. THỐNG KÊ BIẾN ĐỊNH LƯỢNG")
        self.report.append("\n4. THỐNG KÊ BIẾN ĐỊNH LƯỢNG")
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns
        stats_df = self.df[numerical_cols].describe()
        
        # In từng cột một cách rõ ràng để đảm bảo hiển thị đầy đủ
        print(f"\nTổng số biến định lượng: {len(numerical_cols)}")
        print(f"Danh sách biến: {list(numerical_cols)}\n")
        
        # Hiển thị toàn bộ ma trận thống kê
        with pd.option_context('display.max_columns', None, 
                              'display.width', 200,
                              'display.precision', 6):
            print(stats_df)
        
        self.report.append(f"\nTổng số biến định lượng: {len(numerical_cols)}")
        self.report.append(f"Danh sách biến: {list(numerical_cols)}")
        self.report.append("\n" + str(stats_df))
        
        # 5. Categorical statistics
        print("\n5. THỐNG KÊ BIẾN ĐỊNH TÍNH")
        self.report.append("\n5. THỐNG KÊ BIẾN ĐỊNH TÍNH")
        categorical_cols = self.df.select_dtypes(include=['object']).columns
        for col in categorical_cols:
            print(f"\n   {col}:")
            print(self.df[col].value_counts())
            self.report.append(f"\n   {col}:")
            self.report.append(str(self.df[col].value_counts()))
        
        self.analysis_history.append("Phân Tích Thống Kê Mô Tả")
    
    # ============ B) ESTIMATION & HYPOTHESIS TESTING ============
    def estimation_hypothesis_testing(self):
        """
        TÍNH NĂNG: Ước Lượng và Kiểm Định Giả Thuyết
        MÔ TẢ: Thực hiện suy diễn thống kê bao gồm:
        - Khoảng tin cậy 95%
        - Kiểm định T
        - Kiểm định Chi-bình phương
        - Giải thích giá trị P
        """
        print("\n" + "="*80)
        print("B) ƯỚC LƯỢNG VÀ KIỂM ĐỊNH GIẢ THUYẾT")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("B) ƯỚC LƯỢNG VÀ KIỂM ĐỊNH GIẢ THUYẾT")
        self.report.append("="*80)
        
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns
        
        # 1. Confidence intervals for means
        print("\n1. KHOẢNG TIN CẬY 95% CHO TRUNG BÌNH")
        self.report.append("\n1. KHOẢNG TIN CẬY 95% CHO TRUNG BÌNH")
        for col in numerical_cols:
            data = self.df[col].dropna()
            mean = data.mean()
            std = data.std()
            n = len(data)
            se = std / np.sqrt(n)
            ci = stats.t.interval(0.95, n-1, loc=mean, scale=se)
            print(f"\n   {col}:")
            print(f"      Trung bình: {mean:.4f}")
            print(f"      Khoảng tin cậy 95%: [{ci[0]:.4f}, {ci[1]:.4f}]")
            self.report.append(f"\n   {col}: Trung bình={mean:.4f}, Khoảng tin cậy=[{ci[0]:.4f}, {ci[1]:.4f}]")
        
        # 2. T-test: age by deposit status
        print("\n2. KIỂM ĐỊNH T: TUỔI THEO TRẠNG THÁI GỬI TIỀN")
        self.report.append("\n2. KIỂM ĐỊNH T: TUỔI THEO TRẠNG THÁI GỬI TIỀN")
        if 'deposit' in self.df.columns and 'age' in self.df.columns:
            yes_group = self.df[self.df['deposit'] == 'yes']['age']
            no_group = self.df[self.df['deposit'] == 'no']['age']
            t_stat, p_value = stats.ttest_ind(yes_group, no_group)
            print(f"   Nhóm Có (n={len(yes_group)}): Trung bình={yes_group.mean():.4f}, Độ lệch chuẩn={yes_group.std():.4f}")
            print(f"   Nhóm Không (n={len(no_group)}): Trung bình={no_group.mean():.4f}, Độ lệch chuẩn={no_group.std():.4f}")
            print(f"   Thống kê T: {t_stat:.4f}")
            print(f"   Giá trị P: {p_value:.6f}")
            sig = "Có ý nghĩa thống kê" if p_value < 0.05 else "Không có ý nghĩa thống kê"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   Thống kê T={t_stat:.4f}, Giá trị P={p_value:.6f}")
        
        # 3. Chi-square test
        print("\n3. KIỂM ĐỊNH CHI-BÌNH PHƯƠNG: CÔNG VIỆC vs GỬI TIỀN")
        self.report.append("\n3. KIỂM ĐỊNH CHI-BÌNH PHƯƠNG: CÔNG VIỆC vs GỬI TIỀN")
        if 'job' in self.df.columns and 'deposit' in self.df.columns:
            ct = pd.crosstab(self.df['job'], self.df['deposit'])
            chi2, p_value, dof, expected = stats.chi2_contingency(ct)
            print(f"   Chi-bình phương: {chi2:.4f}")
            print(f"   Giá trị P: {p_value:.6f}")
            print(f"   Bậc tự do: {dof}")
            sig = "Có sự liên kết" if p_value < 0.05 else "Không có sự liên kết"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   Chi-bình phương={chi2:.4f}, Giá trị P={p_value:.6f}")
        
        self.analysis_history.append("Ước Lượng & Kiểm Định Giả Thuyết")
    
    # ============ C) CORRELATION ANALYSIS ============
    def correlation_analysis(self):
        """
        TÍNH NĂNG: Phân Tích Tương Quan
        MÔ TẢ: Kiểm tra các mối quan hệ bao gồm:
        - Ma trận tương quan Pearson
        - Tương quan với biến mục tiêu
        - Tương quan hạng Spearman
        """
        print("\n" + "="*80)
        print("C) PHÂN TÍCH TƯƠNG QUAN")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("C) PHÂN TÍCH TƯƠNG QUAN")
        self.report.append("="*80)
        
        numerical_cols = self.df_encoded.select_dtypes(include=[np.number]).columns
        
        # 1. Pearson correlation
        print("\n1. MA TRẬN TƯƠNG QUAN PEARSON")
        self.report.append("\n1. MA TRẬN TƯƠNG QUAN PEARSON")
        corr_matrix = self.df_encoded[numerical_cols].corr()
        print(corr_matrix)
        self.report.append("\n" + str(corr_matrix))
        
        # 2. Correlation with target variable (deposit)
        if 'deposit' in self.df_encoded.columns:
            print("\n2. TƯƠNG QUAN VỚI BIẾN MỤC TIÊU (GỬI TIỀN)")
            self.report.append("\n2. TƯƠNG QUAN VỚI BIẾN MỤC TIÊU (GỬI TIỀN)")
            deposit_corr = corr_matrix['deposit'].sort_values(ascending=False)
            print(deposit_corr)
            self.report.append("\n" + str(deposit_corr))
        
        # 3. Spearman correlation for age vs balance
        print("\n3. TƯƠNG QUAN HẠNG SPEARMAN: TUỔI vs SỐ DƯ")
        self.report.append("\n3. TƯƠNG QUAN HẠNG SPEARMAN: TUỔI vs SỐ DƯ")
        if 'age' in self.df.columns and 'balance' in self.df.columns:
            spearman_corr, p_value = stats.spearmanr(self.df['age'], self.df['balance'])
            print(f"   Hệ số tương quan: {spearman_corr:.4f}")
            print(f"   Giá trị P: {p_value:.6f}")
            self.report.append(f"   Hệ số tương quan: {spearman_corr:.4f}, Giá trị P={p_value:.6f}")
        
        self.analysis_history.append("Phân Tích Tương Quan")
    
    # ============ D) ANOVA ANALYSIS ============
    def anova_analysis(self):
        """
        TÍNH NĂNG: Phân Tích ANOVA
        MÔ TẢ: Kiểm tra sự khác biệt giữa các nhóm:
        - ANOVA một chiều: Tuổi theo công việc
        - ANOVA một chiều: Số dư theo tình trạng hôn nhân
        - Thống kê F và giá trị P
        """
        print("\n" + "="*80)
        print("D) PHÂN TÍCH ANOVA")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("D) PHÂN TÍCH ANOVA")
        self.report.append("="*80)
        
        # 1. One-way ANOVA: Age by job
        print("\n1. ANOVA MỘT CHIỀU: TUỔI THEO LOẠI CÔNG VIỆC")
        self.report.append("\n1. ANOVA MỘT CHIỀU: TUỔI THEO LOẠI CÔNG VIỆC")
        if 'age' in self.df.columns and 'job' in self.df.columns:
            groups = self.df.groupby('job')['age'].apply(list)
            f_stat, p_value = stats.f_oneway(*groups)
            print(f"   Thống kê F: {f_stat:.4f}")
            print(f"   Giá trị P: {p_value:.6f}")
            sig = "Có sự khác biệt đáng kể" if p_value < 0.05 else "Không có sự khác biệt đáng kể"
            print(f"   Kết luận: {sig}")
            self.report.append(f"   Thống kê F={f_stat:.4f}, Giá trị P={p_value:.6f}")
            
            print("\n   Trung bình tuổi theo loại công việc:")
            self.report.append("\n   Trung bình tuổi theo loại công việc:")
            job_means = self.df.groupby('job')['age'].mean().sort_values(ascending=False)
            print(job_means)
            self.report.append("\n" + str(job_means))
        
        # 2. One-way ANOVA: Balance by marital status
        print("\n2. ANOVA MỘT CHIỀU: SỐ DƯ THEO TÌNH TRẠNG HÔN NHÂN")
        self.report.append("\n2. ANOVA MỘT CHIỀU: SỐ DƯ THEO TÌNH TRẠNG HÔN NHÂN")
        if 'balance' in self.df.columns and 'marital' in self.df.columns:
            groups = self.df.groupby('marital')['balance'].apply(list)
            f_stat, p_value = stats.f_oneway(*groups)
            print(f"   Thống kê F: {f_stat:.4f}")
            print(f"   Giá trị P: {p_value:.6f}")
            self.report.append(f"   Thống kê F={f_stat:.4f}, Giá trị P={p_value:.6f}")
        
        self.analysis_history.append("Phân Tích ANOVA")
    
    # ============ E) REGRESSION ANALYSIS ============
    def regression_analysis(self):
        """
        TÍNH NĂNG: Phân Tích Hồi Quy
        MÔ TẢ: Mô hình hóa các mối quan hệ bao gồm:
        - Hồi quy tuyến tính đơn
        - Hồi quy tuyến tính đa biến
        - Giá trị R bình phương
        """
        print("\n" + "="*80)
        print("E) PHÂN TÍCH HỒI QUY")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("E) PHÂN TÍCH HỒI QUY")
        self.report.append("="*80)
        
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numerical_cols) >= 2:
            # 1. Simple linear regression: balance vs age
            print("\n1. HỒI QUY TUYẾN TÍNH ĐƠN: SỐ DƯ vs TUỔI")
            self.report.append("\n1. HỒI QUY TUYẾN TÍNH ĐƠN: SỐ DƯ vs TUỔI")
            if 'age' in self.df.columns and 'balance' in self.df.columns:
                X_simple = self.df[['age']].values
                y_simple = self.df['balance'].values
                model_simple = LinearRegression()
                model_simple.fit(X_simple, y_simple)
                r2_simple = model_simple.score(X_simple, y_simple)
                
                print(f"   Hệ số chặn: {model_simple.intercept_:.4f}")
                print(f"   Hệ số (tuổi): {model_simple.coef_[0]:.4f}")
                print(f"   R bình phương: {r2_simple:.4f}")
                print(f"   Phương trình: Số dư = {model_simple.intercept_:.4f} + {model_simple.coef_[0]:.4f} * Tuổi")
                self.report.append(f"   Hệ số chặn={model_simple.intercept_:.4f}")
                self.report.append(f"   Hệ số={model_simple.coef_[0]:.4f}")
                self.report.append(f"   R bình phương={r2_simple:.4f}")
            
            # 2. Multiple linear regression
            print("\n2. HỒI QUY TUYẾN TÍNH ĐA BIẾN")
            self.report.append("\n2. HỒI QUY TUYẾN TÍNH ĐA BIẾN")
            selected_cols = [col for col in numerical_cols if col != 'balance'][:5]
            if 'balance' in numerical_cols and len(selected_cols) > 0:
                X_multi = self.df[selected_cols].values
                y_multi = self.df['balance'].values
                model_multi = LinearRegression()
                model_multi.fit(X_multi, y_multi)
                r2_multi = model_multi.score(X_multi, y_multi)
                
                print(f"   Biến độc lập: {selected_cols}")
                print(f"   Hệ số chặn: {model_multi.intercept_:.4f}")
                print(f"   Các hệ số:")
                for i, col in enumerate(selected_cols):
                    print(f"      {col}: {model_multi.coef_[i]:.4f}")
                print(f"   R bình phương: {r2_multi:.4f}")
                self.report.append(f"   Biến độc lập: {selected_cols}")
                self.report.append(f"   Hệ số chặn: {model_multi.intercept_:.4f}")
                self.report.append(f"   R bình phương: {r2_multi:.4f}")
        
        self.analysis_history.append("Phân Tích Hồi Quy")
    
    def find_insights(self):
        """
        TÍNH NĂNG: Tự Động Tìm Insights
        MÔ TẢ: Phát hiện các patterns và insights quan trọng từ dữ liệu
        """
        print("\n" + "="*80)
        print("F) TỰ ĐỘNG TÌM INSIGHTS")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("F) TỰ ĐỘNG TÌM INSIGHTS")
        self.report.append("="*80)
        
        insights = []
        
        # 1. Phát hiện outliers
        print("\n1. PHÁT HIỆN OUTLIERS (Giá trị ngoại lai)")
        self.report.append("\n1. PHÁT HIỆN OUTLIERS (Giá trị ngoại lai)")
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns
        
        for col in numerical_cols:
            Q1 = self.df[col].quantile(0.25)
            Q3 = self.df[col].quantile(0.75)
            IQR = Q3 - Q1
            outliers = self.df[(self.df[col] < Q1 - 1.5*IQR) | (self.df[col] > Q3 + 1.5*IQR)]
            
            if len(outliers) > 0:
                pct = (len(outliers) / len(self.df)) * 100
                insight = f"   - {col}: {len(outliers)} outliers ({pct:.2f}%)"
                print(insight)
                self.report.append(insight)
                insights.append(f"Phát hiện {len(outliers)} outliers trong {col}")
        
        # 2. Tương quan mạnh nhất
        print("\n2. CÁC TƯƠNG QUAN MẠNH NHẤT")
        self.report.append("\n2. CÁC TƯƠNG QUAN MẠNH NHẤT")
        
        if hasattr(self, 'df_encoded'):
            corr_matrix = self.df_encoded.select_dtypes(include=[np.number]).corr()
            corr_pairs = []
            
            for i in range(len(corr_matrix.columns)):
                for j in range(i+1, len(corr_matrix.columns)):
                    col1 = corr_matrix.columns[i]
                    col2 = corr_matrix.columns[j]
                    corr_val = corr_matrix.iloc[i, j]
                    
                    if abs(corr_val) > 0.5:
                        corr_pairs.append((col1, col2, corr_val))
            
            corr_pairs.sort(key=lambda x: abs(x[2]), reverse=True)
            
            for col1, col2, corr_val in corr_pairs[:5]:
                direction = "dương" if corr_val > 0 else "âm"
                insight = f"   - {col1} ↔ {col2}: r={corr_val:.4f} (Tương quan {direction} mạnh)"
                print(insight)
                self.report.append(insight)
                insights.append(f"Tương quan mạnh giữa {col1} và {col2}: {corr_val:.4f}")
        
        # 3. Phân phối không chuẩn
        print("\n3. KIỂM TRA PHÂN PHỐI CHUẨN")
        self.report.append("\n3. KIỂM TRA PHÂN PHỐI CHUẨN")
        
        for col in numerical_cols:
            if len(self.df[col]) > 3:
                skewness = stats.skew(self.df[col].dropna())
                kurtosis = stats.kurtosis(self.df[col].dropna())
                
                if abs(skewness) > 1:
                    skew_type = "lệch phải" if skewness > 0 else "lệch trái"
                    insight = f"   - {col}: Skewness={skewness:.4f} ({skew_type} mạnh)"
                    print(insight)
                    self.report.append(insight)
                    insights.append(f"{col} có phân phối {skew_type}")
        
        # 4. Nhóm có trung bình khác biệt lớn
        if 'deposit' in self.df.columns:
            print("\n4. SO SÁNH THEO NHÓM DEPOSIT")
            self.report.append("\n4. SO SÁNH THEO NHÓM DEPOSIT")
            
            for col in numerical_cols:
                yes_mean = self.df[self.df['deposit'] == 'yes'][col].mean()
                no_mean = self.df[self.df['deposit'] == 'no'][col].mean()
                diff_pct = abs((yes_mean - no_mean) / no_mean * 100)
                
                if diff_pct > 20:
                    insight = f"   - {col}: Nhóm Yes={yes_mean:.2f}, Nhóm No={no_mean:.2f} (Chênh lệch {diff_pct:.1f}%)"
                    print(insight)
                    self.report.append(insight)
                    insights.append(f"Chênh lệch lớn về {col} giữa 2 nhóm deposit")
        
        # 5. Biến có ảnh hưởng mạnh đến target
        if 'deposit' in self.df_encoded.columns:
            print("\n5. BIẾN ẢNH HƯỞNG MẠNH ĐẾN DEPOSIT")
            self.report.append("\n5. BIẾN ẢNH HƯỞNG MẠNH ĐẾN DEPOSIT")
            
            deposit_corr = corr_matrix['deposit'].abs().sort_values(ascending=False)
            top_features = deposit_corr[deposit_corr.index != 'deposit'][:5]
            
            for feature, corr_val in top_features.items():
                insight = f"   - {feature}: |r|={corr_val:.4f}"
                print(insight)
                self.report.append(insight)
                insights.append(f"{feature} có ảnh hưởng mạnh đến deposit")
        
        self.insights = insights
        self.analysis_history.append("Tự Động Tìm Insights")
        
        # Tóm tắt insights
        print(f"\n✓ Tổng cộng phát hiện {len(insights)} insights quan trọng")
        self.report.append(f"\n✓ Tổng cộng phát hiện {len(insights)} insights quan trọng")
    
    def create_visualizations(self):
        """
        TÍNH NĂNG: Tạo Trực Quan Hóa
        """
        print("\n" + "="*80)
        print("G) TẠO TRỰC QUAN HÓA DỮ LIỆU")
        print("="*80)
        
        self.report.append("\n" + "="*80)
        self.report.append("G) TẠO TRỰC QUAN HÓA DỮ LIỆU")
        self.report.append("="*80)
        
        numerical_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = self.df.select_dtypes(include=['object']).columns.tolist()
        
        # 1. Phân phối các biến số
        print("\n1. Đang tạo biểu đồ phân phối...")
        fig, axes = plt.subplots(3, 3, figsize=(18, 15))
        fig.suptitle('PHÂN PHỐI CÁC BIẾN ĐỊNH LƯỢNG', fontsize=16, fontweight='bold')
        
        for idx, col in enumerate(numerical_cols[:9]):
            row = idx // 3
            col_idx = idx % 3
            ax = axes[row, col_idx]
            
            ax.hist(self.df[col].dropna(), bins=30, edgecolor='black', alpha=0.7, color='skyblue')
            ax.set_xlabel(col, fontsize=10)
            ax.set_ylabel('Tần số', fontsize=10)
            ax.set_title(f'Phân phối {col}', fontsize=11, fontweight='bold')
            ax.grid(True, alpha=0.3)
        
        # Chú thích cho cả figure: mô tả biến và mục đích
        caption = "Ghi chú: Mỗi histogram hiển thị phân phối giá trị. " \
                  "Mục đích: Hiểu hình dạng phân phối (lệch, đa đỉnh) và giúp phát hiện outliers."
        fig.text(0.5, 0.02, caption, ha='center', fontsize=10)
        
        plt.tight_layout(rect=[0, 0.03, 1, 0.97])
        filename = f'{self.output_dir}/01_phan_phoi_bien_so.png'
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        self.visualizations.append(filename)
        print(f"   ✓ Đã lưu: {filename}")
        
        # 2. Boxplot để phát hiện outliers
        print("\n2. Đang tạo boxplot phát hiện outliers...")
        fig, axes = plt.subplots(2, 4, figsize=(20, 10))
        fig.suptitle('BOXPLOT - PHÁT HIỆN OUTLIERS', fontsize=16, fontweight='bold')
        
        for idx, col in enumerate(numerical_cols[:8]):
            row = idx // 4
            col_idx = idx % 4
            ax = axes[row, col_idx]
            
            ax.boxplot(self.df[col].dropna(), vert=True)
            ax.set_ylabel(col, fontsize=10)
            ax.set_title(f'Boxplot: {col}', fontsize=11, fontweight='bold')
            ax.grid(True, alpha=0.3)
        
        # Chú thích cho cả figure: mô tả biến và mục đích
        caption = "Ghi chú: Boxplot biểu diễn median, IQR và các điểm ngoại lai. " \
                  "Mục đích: Phát hiện outliers và so sánh phân phối theo biến."
        fig.text(0.5, 0.01, caption, ha='center', fontsize=10)
        plt.tight_layout(rect=[0, 0.03, 1, 0.97])
        filename = f'{self.output_dir}/02_boxplot_outliers.png'
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        self.visualizations.append(filename)
        print(f"   ✓ Đã lưu: {filename}")

        # 3. Ma trận tương quan (Heatmap)
        if hasattr(self, 'df_encoded'):
            print("\n3. Đang tạo ma trận tương quan...")
            corr_matrix = self.df_encoded.select_dtypes(include=[np.number]).corr()
            
            plt.figure(figsize=(14, 12))
            sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm', 
                       center=0, square=True, linewidths=1, cbar_kws={"shrink": 0.8})
            plt.title('MA TRẬN TƯƠNG QUAN PEARSON', fontsize=16, fontweight='bold', pad=20)
            plt.tight_layout()
            
            caption = "Ghi chú: Heatmap hiển thị hệ số tương quan Pearson giữa các biến. " \
                      "Mục đích: Xác định các cặp biến có tương quan mạnh (|r| > 0.5) để phân tích sâu hơn."
            plt.gcf().text(0.5, 0.01, caption, ha='center', fontsize=10)
            filename = f'{self.output_dir}/03_ma_tran_tuong_quan.png'
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            self.visualizations.append(filename)
            print(f"   ✓ Đã lưu: {filename}")

        # 4. Scatter plots cho tương quan mạnh
        print("\n4. Đang tạo scatter plots...")
        if len(numerical_cols) >= 2:
            fig, axes = plt.subplots(2, 3, figsize=(18, 12))
            fig.suptitle('SCATTER PLOTS - MỐI QUAN HỆ GIỮA CÁC BIẾN', fontsize=16, fontweight='bold')
            
            plot_idx = 0
            for i in range(len(numerical_cols)):
                for j in range(i+1, len(numerical_cols)):
                    if plot_idx >= 6:
                        break
                    
                    row = plot_idx // 3
                    col_idx = plot_idx % 3
                    ax = axes[row, col_idx]
                    
                    ax.scatter(self.df[numerical_cols[i]], self.df[numerical_cols[j]], 
                             alpha=0.5, s=20, color='steelblue')
                    ax.set_xlabel(numerical_cols[i], fontsize=10)
                    ax.set_ylabel(numerical_cols[j], fontsize=10)
                    
                    # Tính correlation
                    corr = self.df[numerical_cols[i]].corr(self.df[numerical_cols[j]])
                    ax.set_title(f'{numerical_cols[i]} vs {numerical_cols[j]}\nr={corr:.3f}', 
                               fontsize=11, fontweight='bold')
                    ax.grid(True, alpha=0.3)
                    
                    plot_idx += 1
                    
                if plot_idx >= 6:
                    break
            
            # Chú thích cho cả figure: mô tả biến và mục đích
            caption = "Ghi chú: Scatter plot hiển thị mối quan hệ (tuyến tính/không tuyến tính) giữa hai biến. " \
                      "Mục đích: Kiểm tra xu hướng và mật độ điểm (có thể kèm hệ số tương quan r)."
            fig.text(0.5, 0.02, caption, ha='center', fontsize=10)
            plt.tight_layout(rect=[0, 0.03, 1, 0.95])
            filename = f'{self.output_dir}/04_scatter_plots.png'
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            self.visualizations.append(filename)
            print(f"   ✓ Đã lưu: {filename}")

        # 5. So sánh theo nhóm deposit
        if 'deposit' in self.df.columns:
            print("\n5. Đang tạo biểu đồ so sánh theo deposit...")
            fig, axes = plt.subplots(2, 3, figsize=(18, 10))
            fig.suptitle('SO SÁNH CÁC BIẾN THEO DEPOSIT (YES/NO)', fontsize=16, fontweight='bold')
            
            for idx, col in enumerate(numerical_cols[:6]):
                row = idx // 3
                col_idx = idx % 3
                ax = axes[row, col_idx]
                
                yes_data = self.df[self.df['deposit'] == 'yes'][col]
                no_data = self.df[self.df['deposit'] == 'no'][col]
                
                ax.boxplot([yes_data, no_data], labels=['Yes', 'No'])
                ax.set_ylabel(col, fontsize=10)
                ax.set_xlabel('Deposit', fontsize=10)
                ax.set_title(f'{col} theo Deposit', fontsize=11, fontweight='bold')
                ax.grid(True, alpha=0.3)
            
            # Chú thích cho cả figure: mô tả biến và mục đích
            caption = "Ghi chú: So sánh phân phối giữa hai nhóm Deposit (Yes/No). " \
                      "Mục đích: Xem biến số nào khác biệt lớn giữa hai nhóm."
            fig.text(0.5, 0.01, caption, ha='center', fontsize=10)
            filename = f'{self.output_dir}/05_so_sanh_theo_deposit.png'
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            self.visualizations.append(filename)
            print(f"   ✓ Đã lưu: {filename}")

        # 6. Bar charts cho biến phân loại
        print("\n6. Đang tạo bar charts cho biến phân loại...")
        if len(categorical_cols) > 0:
            fig, axes = plt.subplots(2, 2, figsize=(16, 10))
            fig.suptitle('PHÂN PHỐI CÁC BIẾN PHÂN LOẠI', fontsize=16, fontweight='bold')
            
            for idx, col in enumerate(categorical_cols[:4]):
                row = idx // 2
                col_idx = idx % 2
                ax = axes[row, col_idx]
                
                value_counts = self.df[col].value_counts()
                value_counts.plot(kind='bar', ax=ax, color='coral', edgecolor='black')
                ax.set_xlabel(col, fontsize=10)
                ax.set_ylabel('Số lượng', fontsize=10)
                ax.set_title(f'Phân phối {col}', fontsize=11, fontweight='bold')
                ax.grid(True, alpha=0.3, axis='y')
                plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right')
            
            plt.tight_layout()
            filename = f'{self.output_dir}/06_phan_phoi_bien_phan_loai.png'
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            self.visualizations.append(filename)
            print(f"   ✓ Đã lưu: {filename}")
        
        # 7. Pairplot cho top biến quan trọng
        if 'deposit' in self.df.columns and len(numerical_cols) >= 3:
            print("\n7. Đang tạo pairplot...")
            top_cols = numerical_cols[:4] + ['deposit']
            df_sample = self.df[top_cols].sample(min(500, len(self.df)), random_state=42)
            
            pairplot = sns.pairplot(df_sample, hue='deposit', diag_kind='kde', 
                                   palette={'yes': 'green', 'no': 'red'}, 
                                   plot_kws={'alpha': 0.6})
            pairplot.fig.suptitle('PAIRPLOT - MỐI QUAN HỆ ĐA BIẾN', 
                                 fontsize=16, fontweight='bold', y=1.02)
            
            filename = f'{self.output_dir}/07_pairplot.png'
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            self.visualizations.append(filename)
            print(f"   ✓ Đã lưu: {filename}")
        
        # Sau khi tạo xong, lưu mô tả biến
        self.save_variable_descriptions()
        
        print(f"\n✓ Đã tạo {len(self.visualizations)} biểu đồ")
        self.report.append(f"\n✓ Đã tạo {len(self.visualizations)} biểu đồ trong thư mục {self.output_dir}/")
        
        for viz in self.visualizations:
            self.report.append(f"   - {viz}")
        
        self.analysis_history.append("Tạo Trực Quan Hóa")
    
    def save_terminal_output(self):
        """Lưu toàn bộ output terminal"""
        filename = f'{self.output_dir}/terminal_output.txt'
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(self.report))
        print(f"\n✓ Đã lưu terminal output: {filename}")
        return filename
    
    def create_summary_report(self):
        """Tạo báo cáo tổng hợp với insights và visualizations"""
        print("\n" + "="*80)
        print("TẠO BÁO CÁO TỔNG HỢP")
        print("="*80)
        
        summary = []
        summary.append("="*80)
        summary.append("BÁO CÁO PHÂN TÍCH TỔNG HỢP")
        summary.append(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        summary.append("="*80)
        
        summary.append("\n📊 CÁC PHÂN TÍCH ĐÃ THỰC HIỆN:")
        for i, analysis in enumerate(self.analysis_history, 1):
            summary.append(f"   {i}. {analysis}")
        
        if self.insights:
            summary.append("\n💡 CÁC INSIGHTS QUAN TRỌNG:")
            for i, insight in enumerate(self.insights, 1):
                summary.append(f"   {i}. {insight}")
        
        if self.visualizations:
            summary.append("\n📈 CÁC BIỂU ĐỒ ĐÃ TẠO:")
            for i, viz in enumerate(self.visualizations, 1):
                summary.append(f"   {i}. {viz}")
        
        # Thêm file mô tả biến vào summary
        var_desc_file = f'{self.output_dir}/variable_descriptions.txt'
        if os.path.exists(var_desc_file):
            summary.append("\n📝 File mô tả biến:")
            summary.append(f"   - {var_desc_file}")
        
        summary.append("\n" + "="*80)
        summary.append("KẾT THÚC BÁO CÁO")
        summary.append("="*80)
        
        # Lưu summary
        filename = f'{self.output_dir}/00_summary_report.txt'
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
        
        print("\n".join(summary))
        print(f"\n✓ Đã lưu báo cáo tổng hợp: {filename}")
        
        return filename
    
    def run_all_analysis(self):
        """Chạy tất cả các phân tích - Chuẩn hóa PHẢI chạy trước"""
        print("\n" + "="*80)
        print("CHẠY TẤT CẢ PHÂN TÍCH")
        print("="*80)
        print("\n⚠ Lưu ý: Chuẩn hóa dữ liệu sẽ chạy TRƯỚC")
        
        # Normalize data first
        self.normalize_data()
        
        # Then run all analyses
        self.descriptive_statistics()
        self.estimation_hypothesis_testing()
        self.correlation_analysis()
        self.anova_analysis()
        self.regression_analysis()
        
        # NEW: Tìm insights và tạo visualizations
        self.find_insights()
        self.create_visualizations()
        
        # Lưu outputs
        self.save_terminal_output()
        self.create_summary_report()
        
        print("\n" + "="*80)
        print("✓ ĐÃ HOÀN THÀNH TẤT CẢ PHÂN TÍCH!")
        print("="*80)
        print(f"\nKết quả được lưu trong thư mục: {self.output_dir}/")
        print(f"   - {len(self.visualizations)} biểu đồ")
        print(f"   - {len(self.insights)} insights")
        print(f"   - Báo cáo chi tiết")

    def _save_txt(self, filename):
        """Lưu báo cáo dạng TXT"""
        if not filename.endswith('.txt'):
            filename = filename.replace('.docx', '').replace('.pdf', '') + '.txt'
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.report))
            print(f"\n✓ Đã lưu báo cáo TXT vào: {filename}")
        except Exception as e:
            print(f"✗ Lỗi khi lưu file TXT: {e}")
    
    def _save_docx(self, filename):
        """Lưu báo cáo dạng DOCX"""
        if not DOCX_AVAILABLE:
            print("✗ Chưa cài đặt python-docx. Chạy lệnh: pip install python-docx")
            return
        
        if not filename.endswith('.docx'):
            filename = filename.replace('.txt', '').replace('.pdf', '') + '.docx'
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('BÁO CÁO PHÂN TÍCH MARKETING NGÂN HÀNG', 0)
            title.alignment = 1
            
            # Add timestamp
            timestamp = doc.add_paragraph(f"Tạo lúc: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            timestamp.alignment = 1
            
            doc.add_paragraph()
            
            # Add content
            for line in self.report:
                if line.startswith('='):
                    heading = doc.add_heading(line.strip('=').strip(), level=1)
                    heading.alignment = 1
                elif any(line.startswith(x) for x in ['A)', 'B)', 'C)', 'D)', 'E)']):
                    section = doc.add_heading(line, level=2)
                    section.runs[0].font.size = Pt(12)
                elif line.startswith('   '):
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc.save(filename)
            print(f"\n✓ Đã lưu báo cáo DOCX vào: {filename}")
        except Exception as e:
            print(f"✗ Lỗi khi lưu file DOCX: {e}")
    
    def _save_pdf(self, filename):
        """Lưu báo cáo dạng PDF"""
        if not PDF_AVAILABLE:
            print("✗ Chưa cài đặt reportlab. Chạy lệnh: pip install reportlab")
            return
        
        if not filename.endswith('.pdf'):
            filename = filename.replace('.txt', '').replace('.docx', '') + '.pdf'
        
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            
            pdf = SimpleDocTemplate(filename, pagesize=letter,
                                  rightMargin=0.5*inch, leftMargin=0.5*inch,
                                  topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            story = []
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=12,
                spaceBefore=12,
                alignment=TA_LEFT
            )
            
            title = Paragraph("BÁO CÁO PHÂN TÍCH MARKETING NGÂN HÀNG", title_style)
            story.append(title)
            
            timestamp_text = f"Tạo lúc: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
            timestamp = Paragraph(timestamp_text, styles['Normal'])
            story.append(timestamp)
            story.append(Spacer(1, 0.3*inch))
            
            for line in self.report:
                if line.strip() == '':
                    story.append(Spacer(1, 0.1*inch))
                elif line.startswith('='):
                    heading = Paragraph(line.strip('=').strip(), heading_style)
                    story.append(heading)
                else:
                    line_safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p = Paragraph(line_safe, styles['Normal'])
                    story.append(p)
            
            pdf.build(story)
            print(f"\n✓ Đã lưu báo cáo PDF vào: {filename}")
        except Exception as e:
            print(f"✗ Lỗi khi lưu file PDF: {e}")
    
    def _export_menu(self):
        """Hiển thị menu xuất file"""
        self.display_analysis_summary()
        
        print("\n" + "="*80)
        print("XUẤT KẾT QUẢ PHÂN TÍCH")
        print("="*80)
        print("Chọn định dạng xuất file:")
        print("1. TXT (Tệp văn bản)")
        print("2. DOCX (Tài liệu Word)")
        print("3. PDF (Định dạng PDF)")
        print("4. Tất cả định dạng (Xuất tất cả 3 loại)")
        print("0. Quay lại menu chính")
        
        choice = input("\nNhập lựa chọn (0-4): ").strip()
        
        base_filename = 'bank_analysis_report'
        
        if choice == '1':
            self.save_report(f"{base_filename}.txt", 'txt')
        elif choice == '2':
            self.save_report(f"{base_filename}.docx", 'docx')
        elif choice == '3':
            self.save_report(f"{base_filename}.pdf", 'pdf')
        elif choice == '4':
            print("\nĐang xuất sang tất cả các định dạng...")
            self.save_report(f"{base_filename}.txt", 'txt')
            self.save_report(f"{base_filename}.docx", 'docx')
            self.save_report(f"{base_filename}.pdf", 'pdf')
            print("\n✓ Tất cả báo cáo đã được xuất thành công!")
        elif choice == '0':
            pass
        else:
            print(f"Định dạng không được hỗ trợ: {file_format}")
            print("Các định dạng được hỗ trợ: txt, docx, pdf")
    
    def display_analysis_summary(self):
        """Hiển thị tóm tắt các phân tích đã hoàn thành"""
        print("\n" + "="*80)
        print("TÓM TẮT PHÂN TÍCH ĐÃ HOÀN THÀNH")
        print("="*80)
        print(f"\nTổng số phân tích đã chạy: {len(self.analysis_history)}")
        print("\nCác phân tích đã hoàn thành:")
        for i, analysis in enumerate(self.analysis_history, 1):
            print(f"   {i}. {analysis}")
        print("\n" + "="*80)
    
    def save_report(self, filename='bank_analysis_report.txt', file_format='txt'):
        """Lưu báo cáo vào file theo định dạng chỉ định (txt, docx, pdf)"""
        file_format = file_format.lower()
        
        if file_format == 'txt':
            self._save_txt(filename)
        elif file_format == 'docx':
            self._save_docx(filename)
        elif file_format == 'pdf':
            self._save_pdf(filename)
        else:
            print(f"Định dạng không được hỗ trợ: {file_format}")
            print("Các định dạng được hỗ trợ: txt, docx, pdf")
    
    def display_menu(self):
        """Hiển thị menu tương tác"""
        # Check if normalization has been done
        is_normalized = hasattr(self, 'df_normalized')
        
        while True:
            print("\n" + "="*80)
            print("HỆ THỐNG PHÂN TÍCH THỐNG KÊ MARKETING NGÂN HÀNG")
            print("="*80)
            
            # Show normalization status
            norm_status = "✓ ĐÃ CHUẨN HÓA" if is_normalized else "⚠ CHƯA CHUẨN HÓA"
            print(f"\nTrạng thái chuẩn hóa: {norm_status}")
            
            print("\nChọn phân tích cần thực hiện:")
            print("0. Chuẩn Hóa Dữ Liệu - Chuẩn hóa các biến số")
            print("1. Thống Kê Mô Tả - Thông tin cơ bản về dữ liệu")
            print("2. Ước Lượng & Kiểm Định - Khoảng tin cậy và kiểm định")
            print("3. Phân Tích Tương Quan - Mối quan hệ giữa các biến")
            print("4. Phân Tích ANOVA - So sánh trung bình nhóm")
            print("5. Phân Tích Hồi Quy - Mô hình dự đoán")
            print("6. Tự Động Tìm Insights - Phát hiện patterns quan trọng")
            print("7. Tạo Trực Quan Hóa - Tạo tất cả biểu đồ")
            print("8. Chạy Tất Cả Phân Tích - Thực hiện toàn bộ (bao gồm chuẩn hóa, insights, visualizations)")
            print("9. Xuất Kết Quả - Lưu báo cáo")
            print("10. Thoát - Kết thúc chương trình")
            
            choice = input("\nNhập lựa chọn (0-10): ").strip()
            
            if choice == '0':
                self.normalize_data()
                is_normalized = True
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '1':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.descriptive_statistics()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '2':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.estimation_hypothesis_testing()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '3':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.correlation_analysis()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '4':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.anova_analysis()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '5':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.regression_analysis()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '6':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.find_insights()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '7':
                if not is_normalized:
                    print("\n⚠ LỖI: Bạn phải chuẩn hóa dữ liệu trước!")
                    print("   Vui lòng chọn mục 0")
                    input("\nNhấn Enter để tiếp tục...")
                    continue
                self.create_visualizations()
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '8':
                print("\nChạy tất cả các phân tích...")
                self.run_all_analysis()
                is_normalized = True
                input("\nNhấn Enter để tiếp tục...")
            elif choice == '9':
                self._export_menu()
            elif choice == '10':
                print("\nCảm ơn đã sử dụng hệ thống phân tích!")
                break
            else:
                print("\nLựa chọn không hợp lệ. Vui lòng thử lại!")

def main():
    """Hàm chính"""
    filepath = 'bank.csv'
    
    try:
        analysis = BankMarketingAnalysis(filepath)
        analysis.display_menu()
    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file: {filepath}")
    except Exception as e:
        print(f"Lỗi: {e}")

if __name__ == "__main__":
    main()
